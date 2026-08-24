"""Low-level Word field helpers (TOC, PAGE) via raw OOXML.

python-docx has no field API; these helpers build the fldChar/instrText
sequences directly. The v1 PAGE implementation never inserted instrText, so
the footer rendered blank — this module is the fix.
"""
from __future__ import annotations

from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.text.paragraph import Paragraph


def _fld_char(char_type: str):
    from docx.oxml import OxmlElement

    el = OxmlElement("w:fldChar")
    el.set(qn("w:fldCharType"), char_type)
    return el


def _instr_text(instruction: str):
    from docx.oxml import OxmlElement

    el = OxmlElement("w:instrText")
    el.set(qn("xml:space"), "preserve")
    el.text = instruction
    return el


def add_field(paragraph: Paragraph, instruction: str, placeholder: str = "") -> None:
    """Append a Word field (e.g. PAGE, TOC) to a paragraph.

    Word evaluates the field on open/print; the placeholder text is what
    shows before the first field update.
    """
    run = paragraph.add_run()
    run._r.append(_fld_char("begin"))
    run._r.append(_instr_text(instruction))
    run._r.append(_fld_char("separate"))
    text_run = paragraph.add_run(placeholder)
    text_run.font.italic = True
    end_run = paragraph.add_run()
    end_run._r.append(_fld_char("end"))


def add_page_number(paragraph: Paragraph, before: str = "", after: str = "") -> None:
    """Insert 'Page N of M' fields (PAGE + NUMPAGES) into a footer paragraph."""
    if before:
        paragraph.add_run(before)
    add_field(paragraph, r" PAGE \* MERGEFORMAT ", "1")
    paragraph.add_run(" / ")
    add_field(paragraph, r" NUMPAGES \* MERGEFORMAT ", "1")
    if after:
        paragraph.add_run(after)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER


def add_toc(paragraph: Paragraph) -> None:
    """Insert a real TOC field covering heading levels 1-3.

    The user must right-click → 'Update Field' in Word to populate page
    numbers (Word does not evaluate TOC fields on open by default).
    """
    add_field(paragraph, r' TOC \o "1-3" \h \z \u ',
              placeholder="Right-click here and choose 'Update Field' to generate the table of contents.")
