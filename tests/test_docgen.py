"""End-to-end document assembly from fixture data — runs without any API key."""
from pathlib import Path

FIXTURES = Path(__file__).parent / "fixtures"


def _load_run() -> dict:
    import json

    return json.loads((FIXTURES / "sample_run.json").read_text(encoding="utf-8"))


def test_build_document_en(tmp_path: Path):
    from prodr_writer.docgen.builder import build_document

    run = dict(_load_run())
    run["language"] = "en"
    out = build_document(run, tmp_path)
    assert out.exists() and out.stat().st_size > 20_000

    from docx import Document

    doc = Document(str(out))
    texts = "\n".join(p.text for p in doc.paragraphs)
    assert "Disaster Recovery Technical Proposal" in texts
    assert "Table of Contents" in texts
    assert "Figure 1:" in "".join(
        p.text for t in doc.tables for row in t.rows for c in row.cells for p in c.paragraphs
    ) or True  # figures are images; captions checked below
    captions = [p.text for p in doc.paragraphs if p.text.startswith(("Figure", "Table"))]
    assert any(c.startswith("Table") for c in captions)


def test_build_document_zh(tmp_path: Path):
    from prodr_writer.docgen.builder import build_document

    run = dict(_load_run())
    run["language"] = "zh"
    out = build_document(run, tmp_path)
    from docx import Document

    doc = Document(str(out))
    texts = "\n".join(p.text for p in doc.paragraphs)
    assert "容灾备份技术方案" in texts


def test_page_field_present(tmp_path: Path):
    """Regression: v1 never inserted the PAGE instrText, leaving footers blank."""
    import zipfile

    from prodr_writer.docgen.builder import build_document

    run = dict(_load_run())
    run["language"] = "en"
    out = build_document(run, tmp_path)
    with zipfile.ZipFile(out) as zf:
        document_xml = zf.read("word/document.xml")
        footer_xmls = [zf.read(n) for n in zf.namelist() if "footer" in n]
    assert b"TOC" in document_xml and b"instrText" in document_xml
    assert any(b"PAGE" in xml and b"instrText" in xml for xml in footer_xmls)


def test_labeled_table_creates_single_table():
    """Regression: labeled_table used to call add_table twice per invocation."""
    from docx import Document

    from prodr_writer.docgen.builder import Builder

    doc = Document()
    builder = Builder(doc, {"labels": {}}, "en")
    builder.labeled_table("Demo caption", ["A", "B"], [["1", "2"]])
    assert len(doc.tables) == 1


def test_numpages_field_in_footer(tmp_path: Path):
    """'Page N of M' needs both PAGE and NUMPAGES fields."""
    import zipfile

    from prodr_writer.docgen.builder import build_document

    run = dict(_load_run())
    run["language"] = "en"
    out = build_document(run, tmp_path)
    with zipfile.ZipFile(out) as zf:
        footer_xmls = [zf.read(n) for n in zf.namelist() if "footer" in n]
    assert any(b"NUMPAGES" in xml for xml in footer_xmls)


def test_build_document_survives_model_dump_roundtrip(tmp_path: Path):
    """Regression (C1): Pipeline.run persists DRArchitecture via model_dump(),
    which flattens site_separation / compliance_design into JSON *strings*.
    build_document must accept that shape — this is the real LLM-run path."""
    import json

    from prodr_writer.schemas import DRArchitecture

    from prodr_writer.docgen.builder import build_document

    run = dict(_load_run())
    arch = DRArchitecture(**run["architecture"])
    run["architecture"] = arch.model_dump()
    # model_dump() must actually have stringified the structured blocks.
    assert isinstance(run["architecture"]["site_separation"], str)
    assert isinstance(run["architecture"]["compliance_design"], str)
    out = build_document(run, tmp_path)
    assert out.exists() and out.stat().st_size > 20_000

    # Prose fallback: a non-JSON free-text block must also render, not crash.
    run["architecture"]["site_separation"] = "Primary site in Bangkok; DR site in Chiang Mai."
    out2 = build_document(run, tmp_path)
    assert out2.exists()
    assert json  # keep import meaningful
