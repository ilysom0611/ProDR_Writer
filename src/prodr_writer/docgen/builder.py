"""Document assembly: validated pipeline data → professional .docx proposal.

All static strings come from resources/{language}.yaml; nothing about the
client's industry or country is hardcoded here.
"""
from __future__ import annotations

import json
import time
from pathlib import Path
from typing import Dict, List

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

from ..profiles import localized
from .charts import rto_rpo_chart, tier_distribution_chart, topology_diagram
from .fields import add_page_number, add_toc

ACCENT = RGBColor(0x1F, 0x4E, 0x79)   # dark blue for headings
HEADER_BG = "1F4E79"


def _load_resources(language: str) -> Dict:
    path = Path(__file__).parent.parent / "resources" / f"{language}.yaml"
    if not path.exists():
        path = path.with_name("en.yaml")
    import yaml

    return yaml.safe_load(path.read_text(encoding="utf-8"))


def _set_east_asian(style, font_name: str) -> None:
    """Word picks a fallback CJK font unless eastAsia is set explicitly."""
    rpr = style.element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        from docx.oxml import OxmlElement

        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    rfonts.set(qn("w:eastAsia"), font_name)


def _setup_styles(doc: Document, language: str) -> None:
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    if language == "zh":
        _set_east_asian(normal, "宋体")
        normal.font.size = Pt(10.5)
    for level in range(1, 4):
        style = doc.styles[f"Heading {level}"]
        style.font.name = "Calibri"
        style.font.color.rgb = ACCENT
        style.font.bold = True
        sizes = {1: 16, 2: 13, 3: 12}
        style.font.size = Pt(sizes[level])
        if language == "zh":
            _set_east_asian(style, "微软雅黑")


def _shade(cell, hex_color: str) -> None:
    from docx.oxml import OxmlElement

    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), hex_color)
    cell._tc.get_or_add_tcPr().append(shd)


class Builder:
    def __init__(self, doc: Document, res: Dict, language: str):
        self.doc = doc
        self.res = res
        self.lang = language
        self.fig_no = 0
        self.tab_no = 0

    # -- generic helpers -------------------------------------------------
    def heading(self, text: str, level: int = 1) -> None:
        self.doc.add_heading(text, level=level)

    def para(self, text: str, align=None) -> None:
        p = self.doc.add_paragraph(text)
        if align is not None:
            p.alignment = align
        return p

    def bullets(self, items: List[str]) -> None:
        for item in items:
            self.doc.add_paragraph(item, style="List Bullet")

    def labeled_table(self, caption_text: str, headers: List[str], rows: List[List[str]]) -> None:
        """Captioned table (caption above, per international convention)."""
        self.tab_no += 1
        cap = self.doc.add_paragraph()
        cap.paragraph_format.keep_with_next = True
        cap.add_run(f"Table {self.tab_no}: ").bold = True
        cap.add_run(caption_text)
        table = self.doc.add_table(rows=1, cols=len(headers))
        table.style = "Light Grid Accent 1"
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        for i, header in enumerate(headers):
            cell = table.rows[0].cells[i]
            cell.text = header
            _shade(cell, HEADER_BG)
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.bold = True
                    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        for row in rows:
            cells = table.add_row().cells
            for i, value in enumerate(row):
                cells[i].text = "" if value is None else str(value)

    def figure(self, image_path: Path, caption_text: str, width_inches: float = 5.8) -> None:
        self.fig_no += 1
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run().add_picture(str(image_path), width=Inches(width_inches))
        cap = self.doc.add_paragraph()
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap.add_run(f"Figure {self.fig_no}: ").bold = True
        cap.add_run(caption_text)


# ---------------------------------------------------------------------------
# Entry point
# ---------------------------------------------------------------------------

def build_document(run: dict, run_dir: Path) -> Path:
    language = run.get("language", "en")
    res = _load_resources(language)
    inputs = run["input"]
    bia = run["bia"]
    state = run["current_state"]
    strategy = run["strategy"]
    arch = run["architecture"]
    review = run.get("review") or {}
    validation = run.get("validation") or {"findings": [], "passed": True}

    profile_meta = _load_profile_meta(run.get("profile", "generic-enterprise"))
    profile_title = localized(profile_meta.get("title"), language)

    doc = Document()
    _setup_styles(doc, language)

    # Footer with working PAGE field on every section
    section = doc.sections[0]
    section.top_margin = Inches(0.9)
    section.bottom_margin = Inches(0.9)
    footer_para = section.footer.paragraphs[0]
    add_page_number(footer_para, before="", after="")

    b = Builder(doc, res, language)
    labels = res["labels"]

    # -- Cover -----------------------------------------------------------
    for _ in range(5):
        doc.add_paragraph()
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_title = title.add_run(res["cover"]["title"])
    run_title.font.size = Pt(28)
    run_title.font.bold = True
    run_title.font.color.rgb = ACCENT
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_run = subtitle.add_run(
        f"{res['cover']['subtitle_prefix']} {inputs['project_name']}"
    )
    subtitle_run.font.size = Pt(15)
    meta_lines = [
        f"Client: {inputs.get('client_name', '-')}",
        f"Vendor: {inputs.get('vendor_name', '-')}",
        f"Date: {time.strftime('%Y-%m-%d')}",
        f"Profile: {profile_title}",
    ]
    for line in meta_lines:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run(line).font.size = Pt(11)
    doc.add_paragraph()
    conf = doc.add_paragraph()
    conf.alignment = WD_ALIGN_PARAGRAPH.CENTER
    conf_run = conf.add_run(res["cover"]["confidentiality"])
    conf_run.font.size = Pt(9)
    conf_run.font.italic = True
    doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)

    # -- Document Control --------------------------------------------------
    dc = res["document_control"]
    b.heading(dc["heading"], level=1)
    b.labeled_table("", dc["version_table"][0],
                    [["1.0", time.strftime("%Y-%m-%d"), inputs.get("vendor_name", "-"),
                      "Initial issue"]])
    approval_rows = [[role, "", "", ""] for role in dc["roles"]]
    b.labeled_table("", dc["approval_table"][0], approval_rows)
    b.heading(dc["notice_heading"], level=2)
    b.para(dc["notice"].strip())
    doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)

    # -- Table of Contents -------------------------------------------------
    b.heading(res["toc"]["heading"], level=1)
    b.para(res["toc"]["note"])
    toc_para = doc.add_paragraph()
    add_toc(toc_para)
    doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)

    chapters = res["chapters"]

    # -- 1. Executive Summary ---------------------------------------------
    b.heading(chapters["executive_summary"], 1)
    system_count = len(bia.get("business_systems", []))
    b.para(res["sentences"]["exec_summary_intro"].format(
        project=inputs["project_name"],
        vendor=inputs.get("vendor_name", "-"),
        client=inputs.get("client_name", "-"),
        system_count=system_count,
        overall_rto=bia.get("overall_rto", "-"),
        overall_rpo=bia.get("overall_rpo", "-"),
        profile_title=profile_title,
    ))
    if review.get("score") is not None:
        template = (res["sentences"]["review_passed"] if review.get("can_proceed")
                    else res["sentences"]["review_not_passed"])
        b.para(template.format(rounds=run.get("review_rounds", "?"),
                               score=review.get("score")))
    else:
        b.para(res["sentences"]["review_pending"])

    # -- 2. Background ------------------------------------------------------
    b.heading(chapters["background"], 1)
    b.para(res["sentences"]["background_body"].format(
        client=inputs.get("client_name", "-"), vendor=inputs.get("vendor_name", "-")))
    objectives = [
        f"{labels['overall_rto']}: {bia.get('overall_rto', '-')}",
        f"{labels['overall_rpo']}: {bia.get('overall_rpo', '-')}",
        f"{labels['deployment_mode']}: {arch.get('deployment_mode', '-')}",
    ]
    b.bullets(objectives)

    # -- 3. BIA ---------------------------------------------------------------
    b.heading(chapters["bia"], 1)
    b.para(res["sentences"]["bia_intro"])
    b.labeled_table(
        chapters["bia"],
        [labels["system_name"], labels["tier"], labels["rto"], labels["rpo"],
         labels["criticality"], labels["downtime_impact"]],
        [[s["name"], s["tier"], s.get("rto", ""), s.get("rpo", ""),
          s.get("criticality", ""), s.get("max_downtime_impact", "")]
         for s in bia.get("business_systems", [])],
    )
    try:
        fig1 = run_dir / "chart_tier_distribution.png"
        tier_distribution_chart(_dict_to_bia(bia), fig1, res["charts"])
        b.figure(fig1, res["charts"]["chart_tier_title"])
    except Exception as exc:  # noqa: BLE001 — loud but non-fatal
        b.para(f"[Chart unavailable: {exc}]")

    # -- 4. Current State -----------------------------------------------------
    b.heading(chapters["current_state"], 1)
    b.para(res["sentences"]["current_state_intro"])
    infra = state.get("current_infrastructure", {})
    for area_key, value in infra.items():
        if value:
            b.heading(area_key.replace("_", " ").capitalize(), 3)
            b.para(value)
    gap_rows = [[g.get("area", ""), g.get("current_capability", ""),
                 g.get("required_capability", ""), g.get("gap", ""),
                 g.get("risk_level", "")] for g in state.get("gap_analysis", [])]
    if gap_rows:
        b.labeled_table(labels["gap"], [labels["area"], labels["current_capability"],
                                        labels["required_capability"], labels["gap"],
                                        labels["risk_level"]], gap_rows)

    # -- 5. Strategy -----------------------------------------------------------
    b.heading(chapters["strategy"], 1)
    b.para(res["sentences"]["strategy_intro"])
    if strategy.get("overall_strategy"):
        b.para(strategy["overall_strategy"])
    b.labeled_table(labels["protection_mode"],
                    [labels["tier"], labels["protection_mode"], labels["replication"],
                     labels["failover"], labels["rationale"]],
                    [[t["tier"], t.get("protection_mode", ""), t.get("replication", ""),
                      t.get("failover", ""), t.get("rationale", "")]
                     for t in strategy.get("protection_tiers", [])])

    # -- 6. Architecture ---------------------------------------------------------
    b.heading(chapters["architecture"], 1)
    b.para(res["sentences"]["architecture_intro"])
    site_sep = arch.get("site_separation", {}) or {}
    b.labeled_table(labels["deployment_mode"],
                    [labels["deployment_mode"], labels["primary_site"], labels["dr_site"]],
                    [[arch.get("deployment_mode", ""),
                      f"{arch.get('primary_site', {}).get('name', '')} ({arch.get('primary_site', {}).get('location', '')})",
                      f"{arch.get('dr_site', {}).get('name', '')} ({arch.get('dr_site', {}).get('location', '')})"]])
    if site_sep:
        b.para(", ".join(f"{k}: {v}" for k, v in site_sep.items() if v))
    b.labeled_table(labels["recovery_strategy"],
                    [labels["tier"], labels["systems"], labels["recovery_strategy"],
                     labels["rto"], labels["rpo"], labels["replication"], labels["failover"]],
                    [[tier_key, ", ".join(td.get("systems", [])),
                      td.get("recovery_strategy", ""), td.get("rto", ""), td.get("rpo", ""),
                      td.get("replication", ""), td.get("failover", "")]
                     for tier_key, td in sorted(arch.get("tier_definitions", {}).items())])
    for key in ("network_architecture", "storage_architecture",
                "compute_architecture", "failover_automation"):
        if arch.get(key):
            b.heading(key.replace("_", " ").capitalize(), 2)
            b.para(arch[key])
    vendors = arch.get("vendor_recommendations", {}) or {}
    if vendors:
        b.labeled_table("Vendor recommendations",
                        ["Layer", "Recommendation"],
                        [[k.replace("_", " ").capitalize(), v] for k, v in vendors.items() if v])
    try:
        fig2 = run_dir / "chart_topology.png"
        topology_diagram(_dict_to_arch(arch), fig2, res["charts"])
        b.figure(fig2, res["charts"]["chart_replication_link"])
        fig3 = run_dir / "chart_rto_rpo.png"
        rto_rpo_chart(_dict_to_arch(arch), fig3, res["charts"])
        b.figure(fig3, res["charts"]["chart_rto_title"])
    except Exception as exc:  # noqa: BLE001 — loud but non-fatal
        b.para(f"[Chart unavailable: {exc}]")

    # -- 7. Compliance --------------------------------------------------------
    b.heading(chapters["compliance"], 1)
    b.para(res["sentences"]["compliance_intro"])
    comp_rows = []
    for framework in profile_meta.get("compliance_frameworks", []):
        for req in framework.get("requirements", []):
            comp_rows.append([framework.get("name", framework.get("id", "")), req, ""])
    if comp_rows:
        b.labeled_table(labels["requirement"],
                        ["Framework", labels["requirement"], labels["status"]], comp_rows)
    compliance_design = arch.get("compliance_design", {}) or {}
    if any(compliance_design.values()):
        b.labeled_table("Compliance design elements",
                        ["Aspect", "Design"],
                        [[k.replace("_", " ").capitalize(), v] for k, v in compliance_design.items() if v])

    # -- 8. Roadmap ------------------------------------------------------------
    b.heading(chapters["roadmap"], 1)
    b.para(res["sentences"]["roadmap_intro"])
    phases = res["tables"]["roadmap_phases"]
    b.labeled_table(labels["phase"], [labels["phase"], labels["duration"], labels["activities"]],
                    [[p["phase"], p["duration"], p["activities"]] for p in phases])

    # -- 9. Testing & Maintenance ------------------------------------------------
    b.heading(chapters["testing"], 1)
    b.para(res["sentences"]["testing_intro"])
    b.labeled_table(labels["test_type"], [labels["test_type"], labels["frequency"]],
                    [[t["test_type"], t["frequency"]] for t in res["test_programme"]])

    # -- 10. Validation ----------------------------------------------------------
    b.heading(chapters["validation"], 1)
    findings = validation.get("findings", [])
    sev_names = res["severity_names"]
    if not findings:
        b.para(res["sentences"]["validation_passed"])
    else:
        if not validation.get("passed", True):
            b.para(res["sentences"]["validation_failed"])
        b.labeled_table(labels["rule_id"],
                        [labels["rule_id"], labels["severity"], labels["message"]],
                        [[f["rule_id"], sev_names.get(f["severity"], f["severity"]), f["message"]]
                         for f in findings])

    # -- Appendices ----------------------------------------------------------------
    b.heading(chapters["glossary"], 1)
    b.para(res["sentences"]["glossary_note"])
    b.labeled_table("Glossary", ["Term", "Definition"], [list(pair) for pair in res["glossary"]])
    b.heading(chapters["references"], 1)
    b.para(res["sentences"]["references_note"])
    for ref in res["references"]:
        doc.add_paragraph(ref, style="List Number")

    out_path = run_dir / f"{_safe_name(inputs['project_name'])}_{time.strftime('%Y%m%d_%H%M%S')}.docx"
    doc.save(out_path)
    return out_path


def _safe_name(name: str) -> str:
    import re

    slug = re.sub(r'[\\/:*?"<>|]+', "-", name).strip(". ")[:80]
    return re.sub(r"\s+", "-", slug) or "proposal"


def _load_profile_meta(profile_name: str) -> Dict:
    from pathlib import Path as _Path

    base = _Path(__file__).parent.parent / "profiles"
    path = (_Path.home() / ".prodr" / "profiles" / f"{profile_name}.yaml")
    if not path.exists():
        path = base / f"{profile_name}.yaml"
    import yaml

    return yaml.safe_load(path.read_text(encoding="utf-8"))


def _dict_to_bia(data: dict):
    from ..schemas import BIAReport

    return BIAReport.model_validate(data)


def _dict_to_arch(data: dict):
    from ..schemas import DRArchitecture

    return DRArchitecture.model_validate(data)
